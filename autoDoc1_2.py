import tkinter as tk
from tkinter import ttk
import PyPDF2
from docx import Document
import tkinter as tk  # Importe tkinter aqui
from tkinter import Tk, TkVersion, filedialog
from tkinter import simpledialog
from tkinter import messagebox
import os
import re
from docx2pdf import convert
from PIL import Image, ImageTk
from datetime import datetime


# Função para preencher campos em um arquivo PDF
def preencher_pdf(arquivo_pdf, output_folder):
    
    pdf_reader = PyPDF2.PdfFileReader(arquivo_pdf)

    campos_pdf = {}

    # Detectar campos entre {{}}
    for page_num in range(pdf_reader.numPages):
        page = pdf_reader.getPage(page_num)
        page_text = page.extract_text()

        matches = re.finditer(r'{{(.*?)}}', page_text)
        for match in matches:
            campo = simpledialog.askstring('Editar campo', f'Editar campo: {match.group(1)}', parent=window)
            campos_pdf[match.group()] = campo

        page = pdf_reader.getPage(page_num)  # Obter a página original
        pdf_page = PyPDF2.PdfFileWriter()
        pdf_page.addPage(page)

        for campo, valor in campos_pdf.items():
            page_text = page_text.replace(campo, valor)
            pdf_page.getPage(0).mergePage(page)

        # Gerar um novo arquivo PDF com um nome exclusivo na pasta de saída
        base_name, ext = os.path.splitext(arquivo_pdf)
        output_pdf = os.path.join(output_folder, f"{base_name}Editado_{page_num + 1}.pdf")
        with open(output_pdf, 'wb') as output:
            pdf_page.write(output)
        messagebox.showinfo("PDF Salvo", f"O arquivo PDF da página {page_num + 1} foi preenchido e salvo com sucesso!")

# Função para preencher campos em um arquivo do Word
def preencher_word(arquivo_docx, output_folder):
    doc = Document(arquivo_docx)

    campos_docx = {}

    # Detectar campos entre {{}}
    for para in doc.paragraphs:
        text = para.text
        matches = re.finditer(r'{{(.*?)}}', text)
        for match in matches:
            campo = simpledialog.askstring('Editar campo', f'Editar campo: {match.group(1)}', parent=window)
            campos_docx[match.group()] = campo

        for campo, valor in campos_docx.items():
            text = text.replace(campo, valor)
            para.clear()
            para.add_run(text)

    # Gerar um novo arquivo do Word com um nome exclusivo na pasta de saída
    base_name, ext = os.path.splitext(os.path.basename(arquivo_docx))
    output_docx = os.path.join(output_folder, f"{base_name}Editado.docx")
    doc.save(output_docx)
    messagebox.showinfo("Documento do Word Salvo", "O arquivo do Word foi preenchido e salvo com sucesso!")

# Função para converter um arquivo do Word para PDF
def converter_word_para_pdf():
    arquivo_docx = filedialog.askopenfilename(title="Selecione um arquivo do Word", filetypes=[("Documentos do Word", "*.docx")])
    if arquivo_docx:
        try:
            base_name, ext = os.path.splitext(os.path.basename(arquivo_docx))
            output_pdf = os.path.join(output_folder, f"{base_name}.pdf")  # Nome original + .pdf
            convert(arquivo_docx, output_pdf)
            messagebox.showinfo("Documento do Word Convertido para PDF", f"O arquivo do Word foi convertido para PDF e salvo com sucesso em {output_pdf}!")
        except Exception as e:
            messagebox.showerror("Erro na conversão", f"Ocorreu um erro ao converter o arquivo: {str(e)}")


# Função para exibir um guia passo a passo
def exibir_guia():
    guia = """
    --------------------------------------------
    Preparação do arquivo:
    1º - Antes de anexar o arquivo, substitua os campos de 
    preenchimento por Chaves {{ }}
    --------------------------------------------
    Passo 1: Abra o programa.
    Passo 2: Clique no botão "Selecionar Arquivo" para escolher
    o arquivo que deseja preencher.
    Passo 3: Siga as instruções na janela para preencher o 
    arquivo.
    Passo 4: Ao final do preenchimento o arquivo será salvo
    automaticamente em formado WORD(.docx).
    -------------------------------------------
    Bônus: Caso queira, é possível converter qualquer arquivo
    word(.docx) em PDF, basta clicar no botão:

    --------->>>"Converter WORD para PDF"<<<------------
    """
    messagebox.showinfo("Guia Passo a Passo", guia)


def selecionar_arquivo():
    arquivo = filedialog.askopenfilename(title="Selecione um arquivo")
    if arquivo.lower().endswith('.pdf'):
        preencher_pdf(arquivo, output_folder)
    elif arquivo.lower().endswith('.docx'):
        preencher_word(arquivo, output_folder)
    else:
        messagebox.showinfo("Formato não suportado", "Formato de arquivo não suportado.")

# Verificar se a pasta de saída existe, e criá-la se não existir
output_folder = "output_folder"  # Nome da pasta de saída
if not os.path.exists(output_folder):
    os.makedirs(output_folder)

# Criar janela principal
window = tk.Tk()
window.title("AutoDoc - Preenchimento automático de documentos")

# Definir um ícone personalizado
window.iconbitmap("icone_rcode.ico")

# Definir as dimensões da janela
largura_janela = 600
altura_janela = 540

# Obter as dimensões da tela
largura_tela = window.winfo_screenwidth()
altura_tela = window.winfo_screenheight()

# Calcular as coordenadas para centralizar a janela
x_pos = (largura_tela - largura_janela) // 2
y_pos = (altura_tela - altura_janela) // 2

# Configurar a janela com as dimensões desejadas
window.geometry(f"{largura_janela}x{altura_janela}+{x_pos}+{y_pos}")

# Definir um estilo de fonte
font_style = ("Helvetica", 13)  # Pode ajustar o tamanho da fonte aqui

# Converter a imagem JPEG para GIF
jpeg_image = Image.open("RCODEMASTER.png")  # Substitua "rocket.jpg" pelo nome da sua imagem JPEG
gif_image = ImageTk.PhotoImage(jpeg_image)

# Configurar uma imagem de fundo
background_label = tk.Label(window, image=gif_image)
background_label.place(relwidth=1, relheight=1)

# Função para atualizar a lista de arquivos na Treeview (tabela) com datas e horas de criação
def atualizar_lista_arquivos():
    lista_arquivos.delete(*lista_arquivos.get_children())  # Limpar a tabela atual

    # Definir as colunas
    lista_arquivos["columns"] = ("Nome do Arquivo", "Data e Hora de Criação")
    lista_arquivos.column("#0", width=0, stretch=tk.NO)  # Coluna vazia
    lista_arquivos.column("Nome do Arquivo", anchor="w")
    lista_arquivos.column("Data e Hora de Criação", anchor="w")

    # Configurar os cabeçalhos das colunas
    lista_arquivos.heading("#0", text="", anchor="w")
    lista_arquivos.heading("Nome do Arquivo", text="Nome do Arquivo", anchor="w")
    lista_arquivos.heading("Data e Hora de Criação", text="Data e Hora de Criação", anchor="w")

    arquivos_info = []
    for arquivo in os.listdir(output_folder):
        caminho_arquivo = os.path.join(output_folder, arquivo)
        data_hora_criacao = datetime.fromtimestamp(os.path.getctime(caminho_arquivo))
        arquivos_info.append((arquivo, data_hora_criacao))

    # Adicionar os dados à tabela
    for info in sorted(arquivos_info, key=lambda x: x[1]):
        lista_arquivos.insert("", "end", values=info)

    lista_arquivos.bind("<Double-1>", abrir_arquivo)

#evento de abrir arquivo
def abrir_arquivo(event):
    # Obtenha o item selecionado (arquivo) na lista
    selecionado = lista_arquivos.selection()[0]  # Pode lidar com a seleção de múltiplos arquivos se necessário

    # Obtenha o nome do arquivo selecionado
    nome_arquivo = lista_arquivos.item(selecionado, "values")[0]

    # Construa o caminho completo para o arquivo
    caminho_arquivo = os.path.join(output_folder, nome_arquivo)

    # Abra o arquivo com o aplicativo padrão do sistema
    os.system(f'start "" "{caminho_arquivo}"')  # Isso funciona no Windows, mas pode variar de sistema para sistema

# Criar o widget Treeview (tabela) para mostrar a lista de arquivos
lista_arquivos = ttk.Treeview(window, columns=("Nome do Arquivo", "Data e Hora de Criação"), show="headings")
lista_arquivos.pack()

# Botão para atualizar a lista de arquivos
atualizar_lista_button = tk.Button(window, text="Atualizar Lista de Arquivos", command=atualizar_lista_arquivos, font=font_style)
atualizar_lista_button.pack()
# Definir as colunas
lista_arquivos["columns"] = ("Nome do Arquivo", "Data e Hora de Criação")
lista_arquivos.column("#0", width=0, stretch=tk.NO)  # Coluna vazia
lista_arquivos.column("Nome do Arquivo", anchor="w")
lista_arquivos.column("Data e Hora de Criação", anchor="w")

# Configurar os cabeçalhos das colunas
lista_arquivos.heading("#0", text="", anchor="w")
lista_arquivos.heading("Nome do Arquivo", text="Nome do Arquivo", anchor="w")
lista_arquivos.heading("Data e Hora de Criação", text="Data e Hora de Criação", anchor="w")

# Função para atualizar a lista de arquivos na Treeview (tabela) com datas e horas de criação
def atualizar_lista_arquivos():
    lista_arquivos.delete(*lista_arquivos.get_children())  # Limpar a tabela atual

# Botão para selecionar arquivo
selecionar_arquivo_button = tk.Button(window, text="Selecionar Arquivo", command=selecionar_arquivo, font=font_style)
selecionar_arquivo_button.pack()

# Botão para converter Word para PDF
converter_word_para_pdf_button = tk.Button(window, text="Converter Word para PDF", command=converter_word_para_pdf, font=font_style)
converter_word_para_pdf_button.pack()

# Botão para exibir o guia passo a passo
exibir_guia_button = tk.Button(window, text="Exibir Guia Passo a Passo", command=exibir_guia, font=font_style)
exibir_guia_button.pack()



# Iniciar a janela
window.mainloop()